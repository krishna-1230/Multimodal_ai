from flask import Flask, request, send_file, jsonify, send_from_directory
from docx import Document
from htmldocx import HtmlToDocx
from docx.shared import Pt, RGBColor
import io
import uuid
import os
import logging
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Ensure the output directory exists
OUTPUT_DIR = 'output'
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)


@app.route('/convert-docx', methods=['POST'])
def convert_html_to_docx():
    """
    Supports two payload types:
    - HTML input: { "html_content": "<h1>Hello</h1>", "return_file": true, "default_style": {...} }
    - Structured content: { "content": [ {"type":"heading","text":"...","style":{...}}, ... ] }
    Style fields supported: font_name, font_size, bold, italic, color (hex #RRGGBB), alignment (left/center/right)
    """
    try:
        data = request.get_json(force=True)
    except Exception:
        return jsonify({'error': 'Invalid JSON body'}), 400

    if not data:
        return jsonify({'error': 'No JSON data provided'}), 400

    # Optional: request wants file returned directly
    return_file = bool(data.get('return_file', False))

    document = Document()

    def apply_run_style(run, style: dict):
        if not style or not isinstance(style, dict):
            return
        font = run.font
        name = style.get('font_name')
        if name:
            font.name = name
        size = style.get('font_size')
        if size:
            try:
                font.size = Pt(float(size))
            except Exception:
                pass
        if style.get('bold') is not None:
            font.bold = bool(style.get('bold'))
        if style.get('italic') is not None:
            font.italic = bool(style.get('italic'))
        color = style.get('color')
        if color and isinstance(color, str) and color.startswith('#') and len(color) == 7:
            try:
                r = int(color[1:3], 16)
                g = int(color[3:5], 16)
                b = int(color[5:7], 16)
                font.color.rgb = RGBColor(r, g, b)
            except Exception:
                pass

    try:
        # Structured content takes precedence
        if 'content' in data and isinstance(data.get('content'), list):
            for item in data.get('content', []):
                itype = item.get('type')
                style = item.get('style', {})
                if itype == 'heading':
                    p = document.add_heading(level=item.get('level', 1))
                    run = p.add_run(str(item.get('text', '')))
                    apply_run_style(run, style)
                elif itype == 'paragraph':
                    p = document.add_paragraph()
                    run = p.add_run(str(item.get('text', '')))
                    apply_run_style(run, style)
                elif itype == 'list':
                    for li in item.get('items', []):
                        p = document.add_paragraph(style='List Bullet')
                        run = p.add_run(str(li))
                        apply_run_style(run, style)
                elif itype == 'table':
                    rows = max(1, int(item.get('rows', len(item.get('data', [])))))
                    cols = max(1, int(item.get('cols', len(item.get('data', [])[0]) if item.get('data') else 1)))
                    table = document.add_table(rows=rows, cols=cols)
                    for r_idx, row in enumerate(item.get('data', [])):
                        for c_idx, cell_text in enumerate(row):
                            if r_idx < rows and c_idx < cols:
                                cell = table.cell(r_idx, c_idx)
                                cell.text = str(cell_text)
                elif itype == 'image':
                    img_path = item.get('path')
                    width = item.get('width')
                    try:
                        if width:
                            # python-docx expects Inches or Cm; allow inches
                            from docx.shared import Inches
                            document.add_picture(img_path, width=Inches(float(width)))
                        else:
                            document.add_picture(img_path)
                    except Exception as e:
                        logger.warning('Failed to add image to DOCX: %s', e)
                else:
                    # Unknown type: append as paragraph
                    p = document.add_paragraph()
                    run = p.add_run(str(item.get('text', '')))
                    apply_run_style(run, style)

        elif 'html_content' in data:
            html_content = data.get('html_content')
            if not isinstance(html_content, str) or not html_content.strip():
                return jsonify({'error': 'html_content must be a non-empty string'}), 400
            parser = HtmlToDocx()
            parser.add_html_to_document(html_content, document)

            # Apply default_style to all runs if provided
            default_style = data.get('default_style') or {}
            if default_style:
                for para in document.paragraphs:
                    for run in para.runs:
                        apply_run_style(run, default_style)

        else:
            return jsonify({'error': 'Payload must contain either "html_content" or "content"'}), 400

        # Save to in-memory buffer and disk
        bio = io.BytesIO()
        document.save(bio)
        bio.seek(0)

        unique_filename = f"{uuid.uuid4()}.docx"
        output_path = os.path.join(OUTPUT_DIR, unique_filename)
        with open(output_path, 'wb') as f:
            f.write(bio.getvalue())

        download_url = request.url_root.rstrip('/') + '/output/' + unique_filename

        if return_file:
            bio.seek(0)
            logger.info('Returning generated DOCX directly to client')
            return send_file(
                bio,
                as_attachment=True,
                download_name=unique_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        return jsonify({
            "file": {
                "downloadUrl": download_url,
                "name": unique_filename,
                "path": os.path.abspath(output_path)
            },
            "success": True
        })
    except Exception as e:
        logger.exception('Failed to convert to DOCX')
        return jsonify({'error': 'Conversion failed', 'detail': str(e)}), 500


@app.route('/output/<path:filename>')
def download_file(filename):
    # Serve files from the output directory
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)


if __name__ == '__main__':
    app.run(port=3101, debug=True)